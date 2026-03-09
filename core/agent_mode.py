"""
Agent Mode - Code Interpreter Agent
Sinh Python code để thực hiện tasks tự động
"""
import re
from typing import Dict, Any
from core.llm import LLMProvider
from core.executor import CodeExecutor

class AgentMode:
    def __init__(self, llm: LLMProvider):
        self.llm = llm
        self.executor = CodeExecutor()
        self.max_iterations = 10
        self.safe_mode = True  # Ask confirmation for dangerous operations
        
        # Dangerous keywords requiring confirmation
        self.dangerous_keywords = [
            'admin', 'sudo', 'runas', 'elevation',  # Admin rights
            'download', 'requests.get', 'urllib', 'wget', 'curl',  # Downloads
            'rmdir', 'shutil.rmtree', 'os.remove', 'delete',  # Deletion
            'format', 'diskpart',  # Disk operations
            'regedit', 'registry',  # Registry
            'powershell', 'cmd.exe', 'subprocess.call',  # System commands
            'install', 'pip install', 'apt-get',  # Installations
        ]
    
    def execute_task(self, task: str, confirmation_callback=None) -> Dict[str, Any]:
        """
        Execute task autonomously using Code Interpreter approach
        
        Args:
            task: Task description
            confirmation_callback: Function(code, is_dangerous) -> bool
                                  Returns True to execute, False to skip
        
        Returns:
            {
                'success': bool,
                'iterations': int,
                'results': [{'code': str, 'output': str, 'error': str}],
                'final_message': str
            }
        """
        print(f"\n{'='*60}")
        print(f"🎯 TASK: {task}")
        print(f"{'='*60}\n")
        
        results = []
        
        for iteration in range(1, self.max_iterations + 1):
            print(f"\n--- Iteration {iteration}/{self.max_iterations} ---")
            
            # Step 1: LLM generates code
            code = self._generate_code(task, results)
            
            if code is None:
                # Task completed
                return {
                    'success': True,
                    'iterations': iteration - 1,
                    'results': results,
                    'final_message': '✅ Task hoàn thành!'
                }
            
            print(f"\n📝 Generated Code:\n{code}\n")
            
            # Step 2: Check if confirmation needed
            is_dangerous = self._is_dangerous_code(code)
            
            if self.safe_mode and is_dangerous:
                # Dangerous operation - need confirmation
                if confirmation_callback:
                    # Use callback to request confirmation
                    print("⚠️ Dangerous operation detected - requesting confirmation...")
                    user_confirmed = confirmation_callback(code, is_dangerous)
                    
                    if not user_confirmed:
                        print("⏭️ User skipped - continuing...")
                        results.append({
                            'iteration': iteration,
                            'code': code,
                            'output': '',
                            'error': 'User skipped dangerous operation',
                            'skipped': True,
                            'is_dangerous': True
                        })
                        continue
                    else:
                        print("✅ User confirmed - executing...")
                else:
                    # No callback - auto-skip (fallback)
                    print("⚠️ Dangerous operation - AUTO-SKIPPED (no confirmation handler)")
                    results.append({
                        'iteration': iteration,
                        'code': code,
                        'output': '',
                        'error': 'Auto-skipped: Dangerous operation (no confirmation handler)',
                        'skipped': True,
                        'is_dangerous': True
                    })
                    continue
            else:
                print("✅ Safe operation - auto-executing...")
            
            # Step 3: Execute code
            output, error = self.executor.execute(code)
            
            results.append({
                'iteration': iteration,
                'code': code,
                'output': output,
                'error': error
            })
            
            if error:
                print(f"\n❌ Error: {error}")
            else:
                print(f"\n✅ Output:\n{output}")
            
            # Step 4: Check if task completed
            if self._is_task_completed(task, results):
                return {
                    'success': True,
                    'iterations': iteration,
                    'results': results,
                    'final_message': '✅ Task hoàn thành!'
                }
        
        # Max iterations reached
        return {
            'success': False,
            'iterations': self.max_iterations,
            'results': results,
            'final_message': f'⚠️ Đã đạt {self.max_iterations} iterations nhưng chưa hoàn thành'
        }
    
    def _generate_code(self, task: str, previous_results: list) -> str:
        """Generate Python code to execute the task"""
        system_prompt = """Bạn là Liemdai Copilot - Code Interpreter Agent của Liemdai Team.

Khi được hỏi về identity: "Tôi là Liemdai Copilot, được tạo bởi Liemdai Team."

NHIỆM VỤ:
- Phân tích task của user
- Sinh Python code để thực hiện task
- Code phải HOÀN CHỈNH, có thể chạy độc lập
- Dùng libraries phù hợp (docx, PIL, selenium, pyautogui, os, subprocess...)

OUTPUT FORMAT:
Nếu cần sinh code, trả về trong format:
```python
# Code ở đây
```

Nếu task đã hoàn thành, trả về: "TASK_COMPLETED"

VÍ DỤ:

Task: "Tạo Word document 2 trang về AI agents"
Response:
```python
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('AI Agents Overview', 0)

# Trang 1
doc.add_paragraph('''
AI Agents là hệ thống tự động có khả năng perceive, reason, và act trong môi trường...
[Nội dung dài 1 trang]
''')

# Trang 2
doc.add_heading('Applications', level=1)
doc.add_paragraph('''
AI Agents có nhiều ứng dụng như desktop automation, chatbots, game AI...
[Nội dung dài 1 trang]
''')

doc.save('ai_agents.docx')
print("✅ Đã tạo file ai_agents.docx")
```

Task: "Tắt WiFi trên Windows"
Response:
```python
import os
os.system("netsh interface set interface 'Wi-Fi' admin=disable")
print("✅ Đã tắt WiFi")
```

Task: "Mở VSCode và tạo project Python mới"
Response:
```python
import os
import subprocess

# Tạo thư mục project
os.makedirs('my_python_project/src', exist_ok=True)

# Tạo virtual environment
subprocess.run(['python', '-m', 'venv', 'my_python_project/.venv'], check=True)

# Tạo requirements.txt
with open('my_python_project/requirements.txt', 'w') as f:
    f.write('flask\\nrequests\\npandas\\n')

# Tạo main.py
with open('my_python_project/src/main.py', 'w') as f:
    f.write('print("Hello from my project!")')

# Mở VSCode
subprocess.run(['code', 'my_python_project'])
print("✅ Đã tạo project và mở VSCode")
```
"""
        
        # Build prompt with context
        prompt = f"Task: {task}\n\n"
        
        if previous_results:
            prompt += "Previous Attempts:\n"
            for result in previous_results[-3:]:  # Last 3 attempts
                prompt += f"\nIteration {result['iteration']}:\n"
                prompt += f"Code:\n{result['code']}\n"
                if result['error']:
                    prompt += f"Error: {result['error']}\n"
                else:
                    prompt += f"Output: {result['output']}\n"
        
        prompt += "\nGenerate Python code to execute this task:"
        
        # Get LLM response
        response = self.llm.chat(prompt, system_prompt)
        
        # Check if task completed
        if "TASK_COMPLETED" in response:
            return None
        
        # Extract code from response
        code = self._extract_code(response)
        return code
    
    def _extract_code(self, response: str) -> str:
        """Extract Python code from LLM response"""
        # Try to find code block
        pattern = r'```python\n(.*?)\n```'
        matches = re.findall(pattern, response, re.DOTALL)
        
        if matches:
            return matches[0].strip()
        
        # If no code block, return entire response
        return response.strip()
    
    def _is_task_completed(self, task: str, results: list) -> bool:
        """Check if task is completed based on results"""
        if not results:
            return False
        
        last_result = results[-1]
        
        # If last execution had no error and produced output
        if not last_result['error'] and last_result['output']:
            # Simple heuristic: check for success indicators
            output_lower = last_result['output'].lower()
            success_indicators = ['✅', 'done', 'success', 'completed', 'đã', 'thành công']
            
            for indicator in success_indicators:
                if indicator in output_lower:
                    return True
        
        return False
    
    def _is_dangerous_code(self, code: str) -> bool:
        """
        Check if code contains dangerous operations requiring confirmation
        
        Dangerous operations:
        - Admin/sudo commands
        - Downloads from internet
        - File/folder deletion
        - System commands
        - Registry modifications
        - Package installations
        """
        code_lower = code.lower()
        
        for keyword in self.dangerous_keywords:
            if keyword in code_lower:
                return True
        
        return False
